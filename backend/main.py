import io
import json
import os
import tempfile

import anthropic
import openpyxl
from docx import Document as DocxDocument
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from openai import OpenAI
from pydantic import BaseModel
from supabase import create_client, Client

load_dotenv()

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:3001", "https://knowledge-app-six.vercel.app"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

openai_client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
anthropic_client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
sb: Client = create_client(os.environ["SUPABASE_URL"], os.environ["SUPABASE_SERVICE_KEY"])

CHUNK_WORDS = 400
OVERLAP_WORDS = 50


class SetupOrgRequest(BaseModel):
    firm_name: str


class ProcessRequest(BaseModel):
    document_id: str
    organization_id: str
    file_path: str
    source_type: str


class ChatMessage(BaseModel):
    role: str
    content: str


class ChatRequest(BaseModel):
    message: str
    history: list[ChatMessage] = []
    documents_context: list[dict] = []


class ReEmbedRequest(BaseModel):
    document_id: str
    organization_id: str
    transcript: str


def verify_token(authorization: str):
    token = authorization.removeprefix("Bearer ").strip()
    resp = sb.auth.get_user(token)
    if not resp.user:
        raise HTTPException(status_code=401, detail="Unauthorized")
    return resp.user


def chunk_text(text: str) -> list[str]:
    words = text.split()
    chunks, start = [], 0
    while start < len(words):
        chunks.append(" ".join(words[start: start + CHUNK_WORDS]))
        start += CHUNK_WORDS - OVERLAP_WORDS
    return [c for c in chunks if c.strip()]


@app.post("/setup-org")
def setup_org(req: SetupOrgRequest, authorization: str = Header(...)):
    user = verify_token(authorization)
    user_id = str(user.id)

    # Idempotent: return existing org if user already has one
    existing = (
        sb.table("organization_members")
        .select("organization_id")
        .eq("user_id", user_id)
        .execute()
    )
    if existing.data:
        return {"organization_id": existing.data[0]["organization_id"]}

    org = sb.table("organizations").insert({"name": req.firm_name}).execute()
    if not org.data:
        raise HTTPException(status_code=500, detail="Failed to create organization")
    org_id = org.data[0]["id"]

    sb.table("organization_members").insert(
        {"user_id": user_id, "organization_id": org_id, "role": "admin"}
    ).execute()

    return {"organization_id": org_id}


@app.post("/process")
def process_document(req: ProcessRequest, authorization: str = Header(...)):
    verify_token(authorization)

    try:
        ext = req.file_path.rsplit(".", 1)[-1].lower() if "." in req.file_path else "mp3"
        file_bytes: bytes = sb.storage.from_("documents").download(req.file_path)
        docx_path = None

        if ext == "docx":
            doc = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        elif ext == "xlsx":
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            lines = []
            for sheet in wb.worksheets:
                lines.append(f"[{sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(c) for c in row if c is not None)
                    if row_text.strip():
                        lines.append(row_text)
            text = "\n".join(lines)

        else:
            # Audio — transcribe with Whisper
            with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            try:
                with open(tmp_path, "rb") as f:
                    result = openai_client.audio.transcriptions.create(
                        model="whisper-1",
                        file=f,
                        language="it",
                    )
                text = result.text
            finally:
                os.unlink(tmp_path)

            # Generate a .docx transcript file for audio
            doc_record = sb.table("documents").select("title").eq("id", req.document_id).single().execute()
            doc_title = doc_record.data.get("title", "Trascrizione") if doc_record.data else "Trascrizione"

            docx_doc = DocxDocument()
            docx_doc.add_heading(doc_title, level=0)
            for para in text.split("\n"):
                if para.strip():
                    docx_doc.add_paragraph(para)
            buf = io.BytesIO()
            docx_doc.save(buf)

            docx_path = req.file_path.rsplit(".", 1)[0] + "_trascrizione.docx"
            sb.storage.from_("documents").upload(
                path=docx_path,
                file=buf.getvalue(),
                file_options={
                    "content-type": (
                        "application/vnd.openxmlformats-officedocument"
                        ".wordprocessingml.document"
                    )
                },
            )

        # Chunk + embed (common for all types)
        chunks = chunk_text(text)
        rows = []
        for chunk in chunks:
            emb = (
                openai_client.embeddings.create(model="text-embedding-3-small", input=chunk)
                .data[0]
                .embedding
            )
            rows.append({
                "document_id": req.document_id,
                "organization_id": req.organization_id,
                "content": chunk,
                "embedding": emb,
            })

        sb.table("document_chunks").insert(rows).execute()

        update_data: dict = {"status": "ready", "transcript": text}
        if docx_path:
            update_data["docx_path"] = docx_path
        sb.table("documents").update(update_data).eq("id", req.document_id).execute()

        return {"status": "ok", "chunks": len(chunks)}

    except Exception as exc:
        sb.table("documents").update({"status": "error"}).eq("id", req.document_id).execute()
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/re-embed")
def re_embed(req: ReEmbedRequest, authorization: str = Header(...)):
    verify_token(authorization)

    try:
        sb.table("document_chunks").delete().eq("document_id", req.document_id).execute()

        chunks = chunk_text(req.transcript)
        rows = []
        for chunk in chunks:
            emb = (
                openai_client.embeddings.create(model="text-embedding-3-small", input=chunk)
                .data[0]
                .embedding
            )
            rows.append({
                "document_id": req.document_id,
                "organization_id": req.organization_id,
                "content": chunk,
                "embedding": emb,
            })

        sb.table("document_chunks").insert(rows).execute()
        sb.table("documents").update({"transcript": req.transcript}).eq("id", req.document_id).execute()

        return {"status": "ok", "chunks": len(chunks)}

    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/chat")
def chat(req: ChatRequest, authorization: str = Header(...)):
    user = verify_token(authorization)

    membership = (
        sb.table("organization_members")
        .select("organization_id")
        .eq("user_id", str(user.id))
        .single()
        .execute()
    )
    if not membership.data:
        raise HTTPException(status_code=403, detail="Nessuna organizzazione trovata")
    org_id = membership.data["organization_id"]

    q_emb = (
        openai_client.embeddings.create(model="text-embedding-3-small", input=req.message)
        .data[0]
        .embedding
    )

    hits = (
        sb.rpc(
            "match_chunks",
            {"query_embedding": q_emb, "org_id": org_id, "match_count": 5},
        )
        .execute()
        .data
        or []
    )

    context = "\n\n---\n\n".join(h["content"] for h in hits)
    sources = [
        {"document_id": h["document_id"], "content": h["content"][:200]}
        for h in hits
    ]

    docs_list = ""
    if req.documents_context:
        titles = "\n".join(
            f"- {d.get('title') or 'Untitled'} ({d.get('source_type', 'documento')})"
            for d in req.documents_context
        )
        docs_list = f"\n\nDocumenti disponibili nella knowledge base:\n{titles}"

    system = (
        "Sei un assistente AI per professionisti italiani. Rispondi sempre in italiano, "
        "in modo preciso, chiaro e professionale. "
        "Basa le risposte esclusivamente sul contesto fornito dai documenti aziendali. "
        "Se l'informazione richiesta non è presente nel contesto, dichiaralo esplicitamente "
        "senza inventare dati."
        + docs_list +
        "\n\nContesto dai documenti:\n\n" + context
    )

    messages = [{"role": m.role, "content": m.content} for m in req.history]
    messages.append({"role": "user", "content": req.message})

    def generate():
        with anthropic_client.messages.stream(
            model="claude-sonnet-4-6",
            max_tokens=1024,
            system=system,
            messages=messages,
        ) as stream:
            for text in stream.text_stream:
                yield f"data: {json.dumps({'delta': text})}\n\n"
        yield f"data: {json.dumps({'done': True, 'sources': sources})}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )
